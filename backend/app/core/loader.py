"""
Multi-format document loader.
Supports PDF, Markdown, TXT, source code, and web pages.
"""

from __future__ import annotations

from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from pypdf import PdfReader

SUPPORTED_SOURCE_EXTS = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".tsx": "typescript",
    ".jsx": "javascript",
    ".go": "go",
    ".rs": "rust",
    ".java": "java",
    ".cpp": "cpp",
    ".c": "c",
}


def load_file(file_path: Path | str) -> list[Document]:
    """Load a file and return list of Document objects."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in (".docx", ".doc"):
        return _load_docx(path)
    if suffix in (".md", ".markdown"):
        return _load_text(path, "markdown")
    if suffix in SUPPORTED_SOURCE_EXTS:
        return _load_source_code(path, SUPPORTED_SOURCE_EXTS[suffix])
    if suffix in (".txt", ".text"):
        return _load_text(path, "text")
    return _load_text(path, "text")


def load_url(url: str) -> list[Document]:
    """Fetch a web page and return its text content as a Document."""
    response = httpx.get(url, timeout=30.0, follow_redirects=True)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return [Document(page_content=text, metadata={"source": url, "file_type": "web"})]


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(
                Document(
                    page_content=text.strip(),
                    metadata={
                        "source": path.name,
                        "page": i + 1,
                        "file_type": "pdf",
                    },
                )
            )
    return docs


def _load_docx(path: Path) -> list[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    try:
        doc = DocxDocument(str(path))
    except Exception:
        return []

    docs: list[Document] = []
    current_page: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_page:
                docs.append(
                    Document(
                        page_content="\n".join(current_page),
                        metadata={
                            "source": path.name,
                            "file_type": "docx",
                        },
                    )
                )
                current_page = []
            continue
        current_page.append(text)

    if current_page:
        docs.append(
            Document(
                page_content="\n".join(current_page),
                metadata={"source": path.name, "file_type": "docx"},
            )
        )

    return docs


def _load_text(path: Path, file_type: str) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [
        Document(
            page_content=text,
            metadata={"source": path.name, "file_type": file_type},
        )
    ]


def _load_source_code(path: Path, language: str) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [
        Document(
            page_content=text,
            metadata={"source": path.name, "file_type": "code", "language": language},
        )
    ]
